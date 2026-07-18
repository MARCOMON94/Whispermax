from __future__ import annotations

import argparse
import ctypes
import html
import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import webbrowser
from contextlib import asynccontextmanager, contextmanager
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from queue import Queue
from typing import Any
from uuid import uuid4

import uvicorn
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "salidas"
UPLOADS_DIR = DATA_DIR / "videos"
AUDIO_DIR = DATA_DIR / "audio"
TRANSCRIPTIONS_DIR = DATA_DIR / "transcripciones"
MODELS_DIR = DATA_DIR / "modelos"
QUEUE_RESUME_FILE = DATA_DIR / "reanudar_cola.json"
MAX_BATCH_FILES = int(os.environ.get("WHISPERMAX_MAX_BATCH_FILES", "55"))
MAX_UPLOAD_MB = int(os.environ.get("WHISPERMAX_MAX_UPLOAD_MB", "2048"))
MAX_UPLOAD_BYTES = MAX_UPLOAD_MB * 1024 * 1024
UPLOAD_CHUNK_SIZE = 1024 * 1024
FAST_THREADS = int(os.environ.get("WHISPERMAX_FAST_THREADS", str(min(8, os.cpu_count() or 4))))
TRANSCRIBE_CHUNK_SECONDS = int(os.environ.get("WHISPERMAX_CHUNK_SECONDS", "1200"))
TRANSCRIPTION_ENGINE = os.environ.get("WHISPERMAX_ENGINE", "faster").strip().lower()

RESOURCE_PROFILES = {
    "low": {
        "label": "Bajo",
        "threads": 1,
        "ffmpeg_threads": 1,
        "priority": "below_normal",
        "pause_seconds": 1.5,
    },
    "balanced": {
        "label": "Medio",
        "threads": 2,
        "ffmpeg_threads": 1,
        "priority": "below_normal",
        "pause_seconds": 1.0,
    },
    "fast": {
        "label": "Rapido",
        "threads": 4,
        "ffmpeg_threads": 2,
        "priority": "normal",
        "pause_seconds": 0.0,
    },
    "ultrafast": {
        "label": "Ultrarrapido",
        "threads": FAST_THREADS,
        "ffmpeg_threads": 2,
        "priority": "normal",
        "pause_seconds": 0.0,
    },
}
DEFAULT_RESOURCE_MODE = "low"

for thread_var in ("OMP_NUM_THREADS", "MKL_NUM_THREADS", "NUMBA_NUM_THREADS", "OPENBLAS_NUM_THREADS"):
    os.environ.setdefault(thread_var, str(RESOURCE_PROFILES[DEFAULT_RESOURCE_MODE]["threads"]))

ALLOWED_VIDEO_EXTENSIONS = {
    ".avi",
    ".flv",
    ".m4v",
    ".mkv",
    ".mov",
    ".mp4",
    ".mpeg",
    ".mpg",
    ".webm",
    ".wmv",
}


job_queue: Queue[str] = Queue()
jobs: dict[str, "TranscriptionJob"] = {}
jobs_lock = threading.Lock()
worker_lock = threading.Lock()
worker_thread: threading.Thread | None = None
model_cache: dict[str, Any] = {}
model_cache_lock = threading.Lock()


@dataclass
class TranscriptionJob:
    id: str
    original_name: str
    video_path: Path
    model_name: str
    language: str
    resource_mode: str = DEFAULT_RESOURCE_MODE
    include_timestamps: bool = False
    status: str = "En cola"
    created_at: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    started_at: str = ""
    completed_at: str = ""
    docx_name: str = ""
    txt_name: str = ""
    error: str = ""
    detail: str = "Esperando turno"
    progress: int = 0
    updated_at: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    cancel_requested: bool = False
    audio_path: str = ""
    current_process: subprocess.Popen[str] | None = field(default=None, repr=False, compare=False)


class JobCancelled(Exception):
    pass


@asynccontextmanager
async def lifespan(_: FastAPI):
    apply_resource_limits(DEFAULT_RESOURCE_MODE)
    ensure_output_folders()
    ensure_worker_started()
    resume_saved_queue()
    yield


app = FastAPI(title="Whispermax", version="1.0.0", lifespan=lifespan)


def ensure_output_folders() -> None:
    for folder in (UPLOADS_DIR, AUDIO_DIR, TRANSCRIPTIONS_DIR, MODELS_DIR):
        folder.mkdir(parents=True, exist_ok=True)


def get_resource_profile(resource_mode: str) -> dict[str, float | int | str]:
    return RESOURCE_PROFILES.get(resource_mode, RESOURCE_PROFILES[DEFAULT_RESOURCE_MODE])


def set_process_priority(priority: str) -> None:
    if os.name != "nt":
        return

    priority_classes = {
        "idle": 0x00000040,
        "below_normal": 0x00004000,
        "normal": 0x00000020,
    }
    priority_class = priority_classes.get(priority)
    if priority_class is None:
        return

    try:
        handle = ctypes.windll.kernel32.GetCurrentProcess()
        ctypes.windll.kernel32.SetPriorityClass(handle, priority_class)
    except (AttributeError, OSError):
        return


def apply_resource_limits(resource_mode: str) -> None:
    profile = get_resource_profile(resource_mode)
    threads = str(profile["threads"])
    for thread_var in ("OMP_NUM_THREADS", "MKL_NUM_THREADS", "NUMBA_NUM_THREADS", "OPENBLAS_NUM_THREADS"):
        os.environ[thread_var] = threads

    torch = sys.modules.get("torch")
    if torch is not None:
        try:
            torch.set_num_threads(int(profile["threads"]))
            torch.set_num_interop_threads(1)
        except RuntimeError:
            torch.set_num_threads(int(profile["threads"]))
        except AttributeError:
            pass

    set_process_priority(str(profile["priority"]))


def safe_stem(filename: str) -> str:
    stem = Path(filename).stem.strip() or "video"
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem)
    return stem.strip("._-") or "video"


def is_truthy(value: Any) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "yes", "on", "si", "sí"}


def path_inside(path: Path, folder: Path) -> bool:
    try:
        resolved_path = path.resolve()
        resolved_folder = folder.resolve()
    except OSError:
        return False
    return resolved_folder == resolved_path or resolved_folder in resolved_path.parents


def resume_saved_queue() -> None:
    if not QUEUE_RESUME_FILE.exists():
        return

    try:
        payload = json.loads(QUEUE_RESUME_FILE.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError):
        return

    queued = 0
    for item in payload.get("jobs", []):
        video_path = Path(str(item.get("video_path", "")))
        if not path_inside(video_path, UPLOADS_DIR) or not video_path.exists():
            continue
        enqueue_transcription(
            video_path=video_path,
            original_name=str(item.get("original_name") or video_path.name),
            model_name=str(item.get("model_name") or "tiny"),
            language=str(item.get("language") or "es"),
            resource_mode=str(item.get("resource_mode") or DEFAULT_RESOURCE_MODE),
            include_timestamps=is_truthy(item.get("include_timestamps", False)),
        )
        queued += 1

    consumed = QUEUE_RESUME_FILE.with_suffix(".cargada.json")
    try:
        if consumed.exists():
            consumed.unlink()
        QUEUE_RESUME_FILE.replace(consumed)
    except OSError:
        pass


async def save_upload(upload: UploadFile) -> Path:
    extension = Path(upload.filename or "").suffix.lower()
    if extension not in ALLOWED_VIDEO_EXTENSIONS:
        valid = ", ".join(sorted(ALLOWED_VIDEO_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado. Sube un video con extension: {valid}",
        )

    ensure_output_folders()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    destination = UPLOADS_DIR / f"{safe_stem(upload.filename or 'video')}_{timestamp}{extension}"

    total_size = 0
    try:
        with destination.open("wb") as output:
            while chunk := await upload.read(UPLOAD_CHUNK_SIZE):
                total_size += len(chunk)
                if total_size > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail=f"Archivo demasiado grande. Limite actual: {MAX_UPLOAD_MB} MB.",
                    )
                output.write(chunk)
    except Exception:
        delete_if_inside(destination, UPLOADS_DIR)
        raise

    return destination


def require_ffmpeg() -> None:
    if shutil.which("ffmpeg") is None:
        raise RuntimeError(
            "No se encontro ffmpeg. Instala ffmpeg y asegurate de que este disponible en el PATH."
        )


def update_job(
    job: TranscriptionJob | None,
    *,
    status: str | None = None,
    detail: str | None = None,
    progress: int | None = None,
    error: str | None = None,
    completed: bool = False,
) -> None:
    if job is None:
        return

    with jobs_lock:
        if status is not None:
            job.status = status
        if detail is not None:
            job.detail = detail
        if progress is not None:
            job.progress = max(0, min(100, progress))
        if error is not None:
            job.error = error
        job.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if completed:
            job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def is_cancel_requested(job: TranscriptionJob | None) -> bool:
    if job is None:
        return False
    with jobs_lock:
        return job.cancel_requested


def raise_if_cancelled(job: TranscriptionJob | None) -> None:
    if is_cancel_requested(job):
        raise JobCancelled("Cancelado por el usuario.")


def remember_audio_path(job: TranscriptionJob | None, audio_path: Path) -> None:
    if job is None:
        return
    with jobs_lock:
        job.audio_path = str(audio_path)


def set_current_process(job: TranscriptionJob | None, process: subprocess.Popen[str] | None) -> None:
    if job is None:
        return
    with jobs_lock:
        job.current_process = process


def delete_if_inside(path: Path, folder: Path) -> None:
    try:
        resolved_path = path.resolve()
        resolved_folder = folder.resolve()
        if resolved_folder not in resolved_path.parents:
            return
        if resolved_path.exists() and resolved_path.is_file():
            resolved_path.unlink()
    except OSError:
        return


def delete_dir_if_inside(path: Path, folder: Path) -> None:
    try:
        resolved_path = path.resolve()
        resolved_folder = folder.resolve()
        if resolved_folder not in resolved_path.parents:
            return
        if resolved_path.exists() and resolved_path.is_dir():
            shutil.rmtree(resolved_path)
    except OSError:
        return


def cleanup_job_files(job: TranscriptionJob) -> None:
    if job.status != "Error":
        delete_if_inside(job.video_path, UPLOADS_DIR)
    if job.audio_path:
        audio_path = Path(job.audio_path)
        delete_if_inside(audio_path, AUDIO_DIR)
        delete_dir_if_inside(AUDIO_DIR / f"{audio_path.stem}_chunks", AUDIO_DIR)


def cleanup_output_files(job: TranscriptionJob) -> None:
    if job.docx_name:
        delete_if_inside(TRANSCRIPTIONS_DIR / job.docx_name, TRANSCRIPTIONS_DIR)
    if job.txt_name:
        delete_if_inside(TRANSCRIPTIONS_DIR / job.txt_name, TRANSCRIPTIONS_DIR)


def request_job_cancel(job_id: str) -> bool:
    process: subprocess.Popen[str] | None = None
    queued_cancel = False
    with jobs_lock:
        job = jobs.get(job_id)
        if job is None:
            return False
        if job.status in {"Completado", "Cancelado", "Error"}:
            return True

        job.cancel_requested = True
        process = job.current_process
        if job.status == "En cola":
            job.status = "Cancelado"
            job.detail = "Cancelado antes de empezar"
            job.progress = 100
            job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            queued_cancel = True
        else:
            job.status = "Cancelando"
            job.detail = "Cancelando este trabajo"

    if process is not None and process.poll() is None:
        process.terminate()
    if queued_cancel:
        cleanup_job_files(job)
    return True


def extract_audio(video_path: Path, job: TranscriptionJob | None = None) -> Path:
    raise_if_cancelled(job)
    require_ffmpeg()
    update_job(job, status="Extrayendo audio", detail="Convirtiendo video a WAV", progress=15)
    resource_mode = job.resource_mode if job is not None else DEFAULT_RESOURCE_MODE
    ffmpeg_threads = str(get_resource_profile(resource_mode)["ffmpeg_threads"])
    audio_path = AUDIO_DIR / f"{video_path.stem}.wav"
    remember_audio_path(job, audio_path)
    command = [
        "ffmpeg",
        "-y",
        "-threads",
        ffmpeg_threads,
        "-i",
        str(video_path),
        "-vn",
        "-acodec",
        "pcm_s16le",
        "-ar",
        "16000",
        "-ac",
        "1",
        str(audio_path),
    ]

    process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    set_current_process(job, process)
    stderr = ""
    try:
        while True:
            try:
                _, stderr = process.communicate(timeout=0.5)
                break
            except subprocess.TimeoutExpired:
                if is_cancel_requested(job):
                    process.terminate()
                    try:
                        process.communicate(timeout=5)
                    except subprocess.TimeoutExpired:
                        process.kill()
                        process.communicate()
                    raise JobCancelled("Cancelado por el usuario.")
    finally:
        set_current_process(job, None)

    if process.returncode != 0:
        error = stderr.strip() or "ffmpeg no pudo extraer el audio del video."
        raise RuntimeError(error[-1200:])
    raise_if_cancelled(job)
    return audio_path


def load_whisper_model(model_name: str, resource_mode: str = DEFAULT_RESOURCE_MODE):
    apply_resource_limits(resource_mode)
    cache_key = f"openai:{model_name}"
    with model_cache_lock:
        cached_model = model_cache.get(cache_key)
        if cached_model is not None:
            return cached_model

    try:
        import whisper
    except ImportError as exc:
        raise RuntimeError(
            "No se encontro el paquete whisper. Instala las dependencias con: pip install -r requirements.txt"
        ) from exc

    model = whisper.load_model(model_name, download_root=str(MODELS_DIR))
    apply_resource_limits(resource_mode)
    with model_cache_lock:
        model_cache[cache_key] = model
    return model


def load_faster_whisper_model(model_name: str, resource_mode: str = DEFAULT_RESOURCE_MODE):
    profile = get_resource_profile(resource_mode)
    cache_key = f"faster:{model_name}:{profile['threads']}"
    with model_cache_lock:
        cached_model = model_cache.get(cache_key)
        if cached_model is not None:
            return cached_model

    try:
        from faster_whisper import WhisperModel
    except ImportError as exc:
        raise RuntimeError(
            "No se encontro faster-whisper. Instala las dependencias con: pip install -r requirements.txt"
        ) from exc

    model = WhisperModel(
        model_name,
        device="cpu",
        compute_type="int8",
        cpu_threads=int(profile["threads"]),
        download_root=str(MODELS_DIR),
    )
    with model_cache_lock:
        model_cache[cache_key] = model
    return model


@contextmanager
def whisper_progress(
    job: TranscriptionJob | None,
    *,
    start_percent: int = 0,
    end_percent: int = 89,
    label: str = "",
):
    if job is None:
        yield
        return

    import importlib

    transcribe_module = importlib.import_module("whisper.transcribe")
    original_tqdm = transcribe_module.tqdm.tqdm

    class JobProgressBar:
        def __init__(self, *args: Any, **kwargs: Any) -> None:
            self._bar = original_tqdm(*args, **kwargs)

        def __enter__(self):
            self._bar.__enter__()
            return self

        def __exit__(self, exc_type, exc, traceback) -> bool:
            return self._bar.__exit__(exc_type, exc, traceback)

        def update(self, amount: int | float = 1):
            result = self._bar.update(amount)
            total = getattr(self._bar, "total", 0) or 0
            current = getattr(self._bar, "n", 0) or 0
            if total:
                raw_percent = min(99, int((current / total) * 100))
                whisper_percent = min(end_percent, start_percent + int((raw_percent / 100) * (end_percent - start_percent)))
                label_suffix = f" - {label}" if label else ""
                update_job(
                    job,
                    status="Transcribiendo",
                    detail=f"Whisper esta procesando el audio ({whisper_percent}%){label_suffix}",
                    progress=whisper_percent,
                )
            raise_if_cancelled(job)
            return result

        def __getattr__(self, name: str) -> Any:
            return getattr(self._bar, name)

    transcribe_module.tqdm.tqdm = JobProgressBar
    try:
        yield
    finally:
        transcribe_module.tqdm.tqdm = original_tqdm


def media_duration_seconds(path: Path) -> float:
    command = [
        "ffprobe",
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        str(path),
    ]
    try:
        result = subprocess.run(command, capture_output=True, check=True, text=True)
        return max(0.0, float(result.stdout.strip() or 0))
    except (OSError, subprocess.CalledProcessError, ValueError):
        if path.suffix.lower() == ".wav":
            try:
                # 16 kHz, mono, 16-bit PCM = 32,000 bytes per second.
                return max(0.0, path.stat().st_size / 32000)
            except OSError:
                pass
        return 0.0


def split_audio_for_whisper(audio_path: Path, chunk_seconds: int, job: TranscriptionJob | None = None) -> list[Path]:
    chunk_dir = AUDIO_DIR / f"{audio_path.stem}_chunks"
    delete_dir_if_inside(chunk_dir, AUDIO_DIR)
    chunk_dir.mkdir(parents=True, exist_ok=True)
    pattern = chunk_dir / f"{audio_path.stem}_%04d.wav"
    command = [
        "ffmpeg",
        "-y",
        "-i",
        str(audio_path),
        "-f",
        "segment",
        "-segment_time",
        str(chunk_seconds),
        "-reset_timestamps",
        "1",
        "-acodec",
        "pcm_s16le",
        "-ar",
        "16000",
        "-ac",
        "1",
        str(pattern),
    ]

    update_job(job, status="Preparando audio", detail="Dividiendo audio largo en partes", progress=25)
    process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    set_current_process(job, process)
    stderr = ""
    try:
        while True:
            try:
                _, stderr = process.communicate(timeout=0.5)
                break
            except subprocess.TimeoutExpired:
                if is_cancel_requested(job):
                    process.terminate()
                    try:
                        process.communicate(timeout=5)
                    except subprocess.TimeoutExpired:
                        process.kill()
                        process.communicate()
                    raise JobCancelled("Cancelado por el usuario.")
    finally:
        set_current_process(job, None)

    if process.returncode != 0:
        error = stderr.strip() or "ffmpeg no pudo dividir el audio."
        raise RuntimeError(error[-1200:])

    chunks = sorted(chunk_dir.glob("*.wav"))
    if not chunks:
        raise RuntimeError("No se generaron partes de audio para Whisper.")
    return chunks


def merge_transcriptions(parts: list[tuple[dict[str, Any], float]]) -> dict[str, Any]:
    text_parts: list[str] = []
    merged_segments: list[dict[str, Any]] = []
    for transcription, offset in parts:
        text = str(transcription.get("text", "")).strip()
        if text:
            text_parts.append(text)
        for segment in transcription.get("segments", []):
            merged_segment = dict(segment)
            for key in ("start", "end"):
                if key in merged_segment:
                    try:
                        merged_segment[key] = float(merged_segment[key]) + offset
                    except (TypeError, ValueError):
                        pass
            merged_segments.append(merged_segment)

    return {"text": "\n".join(text_parts).strip(), "segments": merged_segments}


def transcribe_with_faster_whisper(
    model: Any,
    audio_path: Path,
    language: str,
    job: TranscriptionJob | None,
    *,
    duration: float,
    start_percent: int,
    end_percent: int,
    label: str = "",
) -> dict[str, Any]:
    options: dict[str, Any] = {
        "beam_size": 1,
        "condition_on_previous_text": False,
        "vad_filter": False,
    }
    if language:
        options["language"] = language

    label_suffix = f" - {label}" if label else ""
    segments_iter, _ = model.transcribe(str(audio_path), **options)
    segments: list[dict[str, Any]] = []
    text_parts: list[str] = []
    for segment in segments_iter:
        raise_if_cancelled(job)
        segment_text = str(segment.text or "").strip()
        if segment_text:
            text_parts.append(segment_text)
        segments.append(
            {
                "id": len(segments),
                "start": float(segment.start),
                "end": float(segment.end),
                "text": segment_text,
            }
        )
        if duration > 0:
            chunk_ratio = min(1.0, max(0.0, float(segment.end) / duration))
            progress = min(end_percent, start_percent + int(chunk_ratio * (end_percent - start_percent)))
            update_job(
                job,
                status="Transcribiendo",
                detail=f"Faster Whisper esta procesando el audio ({progress}%){label_suffix}",
                progress=progress,
            )

    update_job(
        job,
        status="Transcribiendo",
        detail=f"Faster Whisper esta procesando el audio ({end_percent}%){label_suffix}",
        progress=end_percent,
    )
    return {"text": " ".join(text_parts).strip(), "segments": segments}


def transcribe_audio(
    audio_path: Path,
    model_name: str,
    language: str,
    job: TranscriptionJob | None = None,
) -> dict[str, Any]:
    resource_mode = job.resource_mode if job is not None else DEFAULT_RESOURCE_MODE
    raise_if_cancelled(job)
    use_faster = TRANSCRIPTION_ENGINE != "openai"
    engine_label = "Faster Whisper" if use_faster else "Whisper"
    update_job(job, status="Cargando modelo", detail=f"Preparando {engine_label} {model_name}", progress=30)
    try:
        model = load_faster_whisper_model(model_name, resource_mode=resource_mode) if use_faster else load_whisper_model(model_name, resource_mode=resource_mode)
    except Exception:
        if not use_faster:
            raise
        use_faster = False
        engine_label = "Whisper"
        update_job(job, status="Cargando modelo", detail=f"Preparando Whisper {model_name}", progress=30)
        model = load_whisper_model(model_name, resource_mode=resource_mode)
    apply_resource_limits(resource_mode)
    raise_if_cancelled(job)
    profile = get_resource_profile(resource_mode)
    update_job(
        job,
        status="Transcribiendo",
        detail=f"{engine_label} esta procesando el audio (0%) con consumo {profile['label']} ({profile['threads']} hilo/s)",
        progress=0,
    )

    duration = media_duration_seconds(audio_path)
    if duration > TRANSCRIBE_CHUNK_SECONDS:
        chunks = split_audio_for_whisper(audio_path, TRANSCRIBE_CHUNK_SECONDS, job=job)
        parts: list[tuple[dict[str, Any], float]] = []
        total_chunks = len(chunks)
        try:
            for index, chunk_path in enumerate(chunks, start=1):
                raise_if_cancelled(job)
                start_percent = int(((index - 1) / total_chunks) * 89)
                end_percent = max(start_percent + 1, int((index / total_chunks) * 89))
                label = f"parte {index}/{total_chunks}"
                update_job(
                    job,
                    status="Transcribiendo",
                    detail=f"{engine_label} esta procesando el audio ({start_percent}%) - {label}",
                    progress=start_percent,
                )
                chunk_duration = media_duration_seconds(chunk_path) or float(TRANSCRIBE_CHUNK_SECONDS)
                if use_faster:
                    chunk_transcription = transcribe_with_faster_whisper(
                        model,
                        chunk_path,
                        language,
                        job,
                        duration=chunk_duration,
                        start_percent=start_percent,
                        end_percent=end_percent,
                        label=label,
                    )
                else:
                    options: dict[str, Any] = {"fp16": False, "verbose": False}
                    if language:
                        options["language"] = language
                    with whisper_progress(
                        job,
                        start_percent=start_percent,
                        end_percent=end_percent,
                        label=label,
                    ):
                        chunk_transcription = model.transcribe(str(chunk_path), **options)
                parts.append((chunk_transcription, float((index - 1) * TRANSCRIBE_CHUNK_SECONDS)))
        finally:
            delete_dir_if_inside(AUDIO_DIR / f"{audio_path.stem}_chunks", AUDIO_DIR)
        transcription = merge_transcriptions(parts)
    else:
        if use_faster:
            transcription = transcribe_with_faster_whisper(
                model,
                audio_path,
                language,
                job,
                duration=duration,
                start_percent=0,
                end_percent=89,
            )
        else:
            options = {"fp16": False, "verbose": False}
            if language:
                options["language"] = language
            with whisper_progress(job):
                transcription = model.transcribe(str(audio_path), **options)
    raise_if_cancelled(job)
    update_job(job, status="Generando DOCX", detail="Creando documentos de salida", progress=90)
    return transcription


def write_docx(
    *,
    source_video: Path,
    audio_path: Path,
    transcription: dict[str, Any],
    model_name: str,
    include_timestamps: bool = False,
    job: TranscriptionJob | None = None,
) -> tuple[Path, Path]:
    raise_if_cancelled(job)
    ensure_output_folders()
    text = str(transcription.get("text", "")).strip()
    segments = transcription.get("segments", [])
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    timestamped_lines = build_timestamped_lines(segments)

    txt_path = TRANSCRIPTIONS_DIR / f"{source_video.stem}_transcripcion.txt"
    docx_path = TRANSCRIPTIONS_DIR / f"{source_video.stem}_transcripcion.docx"

    if include_timestamps and timestamped_lines:
        txt_path.write_text("\n".join(timestamped_lines) + "\n", encoding="utf-8")
    else:
        txt_path.write_text(text + "\n", encoding="utf-8")

    document = Document()
    document.add_heading("Transcripcion de video", level=1)
    document.add_paragraph(f"Archivo original: {source_video.name}")
    document.add_paragraph(f"Audio extraido: {audio_path.name}")
    document.add_paragraph(f"Modelo Whisper: {model_name}")
    document.add_paragraph(f"Fecha: {timestamp}")

    if include_timestamps and timestamped_lines:
        document.add_heading("Texto con marcas de tiempo", level=2)
        for line in timestamped_lines:
            document.add_paragraph(line)
    else:
        document.add_heading("Texto completo", level=2)
        document.add_paragraph(text or "(Sin texto detectado)")

    document.save(docx_path)
    raise_if_cancelled(job)
    return docx_path, txt_path


def build_timestamped_lines(segments: Any) -> list[str]:
    if not isinstance(segments, list):
        return []

    lines = []
    for segment in segments:
        if not isinstance(segment, dict):
            continue
        segment_text = str(segment.get("text", "")).strip()
        if not segment_text:
            continue
        start = format_seconds(segment.get("start", 0))
        end = format_seconds(segment.get("end", 0))
        lines.append(f"[{start} - {end}] {segment_text}")
    return lines


def format_seconds(value: Any) -> str:
    try:
        total = int(float(value))
    except (TypeError, ValueError):
        total = 0
    hours, remainder = divmod(total, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def transcribe_video(
    video_path: Path,
    model_name: str,
    language: str,
    include_timestamps: bool = False,
    job: TranscriptionJob | None = None,
) -> tuple[Path, Path]:
    audio_path = extract_audio(video_path, job=job)
    transcription = transcribe_audio(audio_path, model_name=model_name, language=language, job=job)
    return write_docx(
        source_video=video_path,
        audio_path=audio_path,
        transcription=transcription,
        model_name=model_name,
        include_timestamps=include_timestamps,
        job=job,
    )


def ensure_worker_started() -> None:
    global worker_thread
    with worker_lock:
        if worker_thread is not None and worker_thread.is_alive():
            return
        worker_thread = threading.Thread(target=queue_worker, daemon=True)
        worker_thread.start()


def queue_worker() -> None:
    while True:
        job_id = job_queue.get()
        try:
            with jobs_lock:
                job = jobs.get(job_id)
                if job is None:
                    continue
                if job.cancel_requested or job.status == "Cancelado":
                    job.status = "Cancelado"
                    job.detail = "Cancelado antes de empezar"
                    job.progress = 100
                    if not job.completed_at:
                        job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    continue
                job.status = "Procesando"
                job.detail = "Preparando trabajo"
                job.progress = 5
                job.started_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                resource_mode = job.resource_mode

            try:
                apply_resource_limits(resource_mode)
                docx_path, txt_path = transcribe_video(
                    job.video_path,
                    job.model_name,
                    job.language,
                    include_timestamps=job.include_timestamps,
                    job=job,
                )
            except JobCancelled as exc:
                with jobs_lock:
                    job.status = "Cancelado"
                    job.detail = str(exc)
                    job.error = ""
                    job.progress = 100
                    job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    cleanup_output_files(job)
            except Exception as exc:
                with jobs_lock:
                    job.status = "Error"
                    job.detail = "No se pudo transcribir"
                    job.progress = 100
                    job.error = str(exc)
                    job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            else:
                with jobs_lock:
                    job.status = "Completado"
                    job.detail = "Transcripcion lista"
                    job.progress = 100
                    job.docx_name = docx_path.name
                    job.txt_name = txt_path.name
                    job.completed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        finally:
            if "job" in locals() and job is not None:
                cleanup_job_files(job)
                pause_seconds = float(get_resource_profile(job.resource_mode)["pause_seconds"])
                if pause_seconds:
                    time.sleep(pause_seconds)
            job_queue.task_done()


def enqueue_transcription(
    *,
    video_path: Path,
    original_name: str,
    model_name: str,
    language: str,
    resource_mode: str,
    include_timestamps: bool = False,
) -> TranscriptionJob:
    job = TranscriptionJob(
        id=uuid4().hex,
        original_name=original_name,
        video_path=video_path,
        model_name=model_name,
        language=language,
        resource_mode=resource_mode,
        include_timestamps=include_timestamps,
    )
    with jobs_lock:
        jobs[job.id] = job
    job_queue.put(job.id)
    return job


def job_to_dict(job: TranscriptionJob) -> dict[str, str]:
    progress = str(job.progress)
    if job.status == "Transcribiendo":
        match = re.search(r"\((\d{1,3})%\)", job.detail)
        if match:
            progress = str(max(0, min(100, int(match.group(1)))))

    return {
        "id": job.id,
        "original_name": job.original_name,
        "model_name": job.model_name,
        "language": job.language,
        "resource_mode": job.resource_mode,
        "resource_label": str(get_resource_profile(job.resource_mode)["label"]),
        "include_timestamps": str(job.include_timestamps).lower(),
        "status": job.status,
        "detail": job.detail,
        "progress": progress,
        "created_at": job.created_at,
        "started_at": job.started_at,
        "updated_at": job.updated_at,
        "completed_at": job.completed_at,
        "docx_name": job.docx_name,
        "txt_name": job.txt_name,
        "error": job.error,
        "can_cancel": str(job.status not in {"Completado", "Cancelado", "Error", "Cancelando"}).lower(),
    }


@app.get("/", response_class=HTMLResponse)
async def home() -> str:
    return HTML_PAGE


@app.get("/transcribir", response_class=HTMLResponse)
async def current_queue_page() -> str:
    with jobs_lock:
        current_jobs = list(jobs.values())
    if not current_jobs:
        return HTML_PAGE
    return queue_page(current_jobs)


@app.post("/transcribir", response_class=HTMLResponse)
async def transcribe(
    videos: list[UploadFile] = File(...),
    modelo: str = Form("tiny"),
    idioma: str = Form("es"),
    consumo: str = Form(DEFAULT_RESOURCE_MODE),
    marcas_tiempo: str | None = Form(None),
) -> str:
    model_name = modelo.strip() or "tiny"
    language = idioma.strip()
    resource_mode = consumo.strip() or DEFAULT_RESOURCE_MODE
    if resource_mode not in RESOURCE_PROFILES:
        resource_mode = DEFAULT_RESOURCE_MODE
    include_timestamps = is_truthy(marcas_tiempo)
    ensure_worker_started()

    try:
        if len(videos) > MAX_BATCH_FILES:
            raise HTTPException(
                status_code=400,
                detail=f"Demasiados videos en una tanda. Limite actual: {MAX_BATCH_FILES}.",
            )
        submitted_jobs = []
        for video in videos:
            video_path = await save_upload(video)
            submitted_jobs.append(
                enqueue_transcription(
                    video_path=video_path,
                    original_name=video.filename or video_path.name,
                    model_name=model_name,
                    language=language,
                    resource_mode=resource_mode,
                    include_timestamps=include_timestamps,
                )
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    return queue_page(submitted_jobs)


@app.get("/estado")
async def queue_status(ids: str = "") -> dict[str, list[dict[str, str]]]:
    requested_ids = [job_id.strip() for job_id in ids.split(",") if job_id.strip()]
    with jobs_lock:
        if requested_ids:
            selected_jobs = [jobs[job_id] for job_id in requested_ids if job_id in jobs]
        else:
            selected_jobs = list(jobs.values())
        return {"jobs": [job_to_dict(job) for job in selected_jobs]}


@app.post("/cancelar/{job_id}")
async def cancel_job(job_id: str) -> dict[str, str]:
    if not request_job_cancel(job_id):
        raise HTTPException(status_code=404, detail="Trabajo no encontrado")
    return {"status": "ok"}


@app.get("/descargar/{filename}")
async def download(filename: str) -> FileResponse:
    path = (TRANSCRIPTIONS_DIR / filename).resolve()
    if TRANSCRIPTIONS_DIR.resolve() not in path.parents or not path.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(path, filename=path.name)


def queue_page(submitted_jobs: list[TranscriptionJob]) -> str:
    job_ids = json.dumps([job.id for job in submitted_jobs])
    return f"""
    <!doctype html>
    <html lang="es">
      <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <title>Whispermax</title>
        <style>{CSS}</style>
      </head>
      <body>
        <main class="panel wide">
          <h1>Cola de transcripcion</h1>
          <p>Se procesa un solo video cada vez. Los documentos terminados quedan en <code>salidas/transcripciones</code>.</p>

          <section class="active-job" id="active-job">
            <div>
              <span class="eyebrow">Ahora</span>
              <h2 id="active-title">Preparando cola</h2>
              <p id="active-detail">Leyendo estado...</p>
            </div>
            <div class="progress large">
              <div class="progress-fill" id="active-progress-fill" style="width: 0%"></div>
            </div>
            <div class="active-footer">
              <span class="progress-text" id="active-progress-text">0%</span>
              <div class="actions inline" id="active-actions"></div>
            </div>
          </section>

          <details class="queue-section" open>
            <summary><span id="pending-count">0</span> pendientes</summary>
            <ul class="job-list" id="pending-list"></ul>
          </details>

          <details class="queue-section" id="done-section">
            <summary><span id="done-count">0</span> terminados</summary>
            <ul class="job-list" id="done-list"></ul>
          </details>

          <a class="back" href="/">Anadir mas videos</a>
        </main>
        <script>
          const jobIds = {job_ids};

          function downloadLinks(job) {{
            if (!job.docx_name || !job.txt_name) {{
              return "";
            }}
            const docx = encodeURIComponent(job.docx_name);
            const txt = encodeURIComponent(job.txt_name);
            return `<a class="button small" href="/descargar/${{docx}}">DOCX</a><a class="button secondary small" href="/descargar/${{txt}}">TXT</a>`;
          }}

          function isTerminal(job) {{
            return job.status === "Completado" || job.status === "Cancelado" || job.status === "Error";
          }}

          function isActive(job) {{
            return !isTerminal(job) && job.status !== "En cola";
          }}

          function escapeHtml(value) {{
            return String(value ?? "").replace(/[&<>"']/g, (char) => ({{
              "&": "&amp;",
              "<": "&lt;",
              ">": "&gt;",
              "\\"": "&quot;",
              "'": "&#039;"
            }}[char]));
          }}

          function compactName(name) {{
            return escapeHtml(name || "Video");
          }}

          function progressFromJob(job) {{
            const detailMatch = String(job.detail || "").match(/\\((\\d{{1,3}})%\\)/);
            if (job.status === "Transcribiendo" && detailMatch) {{
              return Math.max(0, Math.min(100, Number(detailMatch[1])));
            }}
            return Math.max(0, Math.min(100, Number(job.progress || 0)));
          }}

          async function cancelJob(jobId) {{
            const button = document.querySelector(`[data-cancel="${{jobId}}"]`);
            if (button) {{
              button.disabled = true;
              button.textContent = "Cancelando";
            }}
            await fetch(`/cancelar/${{jobId}}`, {{ method: "POST", cache: "no-store" }});
            refreshQueue();
          }}

          function renderActive(job) {{
            const wrapper = document.getElementById("active-job");
            const title = document.getElementById("active-title");
            const detail = document.getElementById("active-detail");
            const fill = document.getElementById("active-progress-fill");
            const text = document.getElementById("active-progress-text");
            const actions = document.getElementById("active-actions");

            if (!job) {{
              wrapper.classList.remove("working");
              wrapper.classList.remove("waiting-progress");
              title.textContent = "Sin video activo";
              detail.textContent = "La cola esta esperando el siguiente trabajo.";
              fill.style.width = "0%";
              text.textContent = "0%";
              actions.innerHTML = "";
              return;
            }}

            const progress = progressFromJob(job);
            const isTranscribing = job.status === "Transcribiendo";
            const isWaitingForFirstProgress = isTranscribing && progress === 0;
            wrapper.classList.toggle("working", isTranscribing);
            wrapper.classList.toggle("waiting-progress", isWaitingForFirstProgress);
            title.textContent = job.original_name || "Video";
            detail.textContent = `${{job.status}} · ${{job.error || job.detail || job.model_name}}`;
            fill.style.width = isWaitingForFirstProgress ? "18%" : `${{progress}}%`;
            text.textContent = isWaitingForFirstProgress ? `${{progress}}% real - trabajando` : `${{progress}}%`;
            actions.innerHTML = `${{downloadLinks(job)}}<button class="button danger small" type="button" data-cancel="${{job.id}}" onclick="cancelJob('${{job.id}}')" ${{job.can_cancel !== "true" ? "disabled" : ""}}>Cancelar</button>`;
          }}

          function renderPending(jobs) {{
            const list = document.getElementById("pending-list");
            document.getElementById("pending-count").textContent = jobs.length;
            list.innerHTML = jobs.map((job) => `
              <li class="job-row">
                <span>${{compactName(job.original_name)}}</span>
                <button class="button danger small" type="button" data-cancel="${{job.id}}" onclick="cancelJob('${{job.id}}')" ${{job.can_cancel !== "true" ? "disabled" : ""}}>Cancelar</button>
              </li>
            `).join("") || `<li class="empty-row">No hay videos pendientes.</li>`;
          }}

          function renderDone(jobs) {{
            const list = document.getElementById("done-list");
            document.getElementById("done-count").textContent = jobs.length;
            list.innerHTML = jobs.map((job) => `
              <li class="job-row done">
                <span>${{compactName(job.original_name)}}<small>${{escapeHtml(job.status)}}</small></span>
                <span class="download-cell">${{downloadLinks(job)}}</span>
              </li>
            `).join("") || `<li class="empty-row">Todavia no hay trabajos terminados en esta cola.</li>`;
          }}

          async function refreshQueue() {{
            const response = await fetch(`/estado?ids=${{jobIds.join(",")}}`, {{ cache: "no-store" }});
            const data = await response.json();
            const active = data.jobs.find(isActive);
            const pendingJobs = data.jobs.filter((job) => job.status === "En cola");
            const doneJobs = data.jobs.filter(isTerminal);
            const stillRunning = data.jobs.some((job) => !isTerminal(job));

            renderActive(active);
            renderPending(pendingJobs);
            renderDone(doneJobs);

            if (stillRunning) {{
              window.setTimeout(refreshQueue, 2500);
            }}
          }}

          refreshQueue();
        </script>
      </body>
    </html>
    """


CSS = """
:root {
  color-scheme: light;
  font-family: Arial, Helvetica, sans-serif;
  background: #f4f6f8;
  color: #1f2933;
}

* {
  box-sizing: border-box;
}

body {
  align-items: center;
  display: flex;
  justify-content: center;
  margin: 0;
  min-height: 100vh;
  padding: 24px;
}

.panel {
  background: #ffffff;
  border: 1px solid #d9e2ec;
  border-radius: 8px;
  box-shadow: 0 18px 50px rgba(31, 41, 51, 0.08);
  max-width: 560px;
  padding: 28px;
  width: 100%;
}

.panel.wide {
  max-width: 920px;
}

h1 {
  font-size: 28px;
  line-height: 1.15;
  margin: 0 0 10px;
}

p {
  color: #52606d;
  line-height: 1.5;
  margin: 0 0 24px;
}

label {
  display: block;
  font-weight: 700;
  margin: 18px 0 8px;
}

input,
select {
  border: 1px solid #bcccdc;
  border-radius: 6px;
  font: inherit;
  padding: 10px;
  width: 100%;
}

input[type="checkbox"] {
  width: auto;
}

.option-row {
  align-items: center;
  display: flex;
  gap: 10px;
  margin-top: 18px;
}

.option-row label {
  margin: 0;
}

.row {
  display: grid;
  gap: 12px;
  grid-template-columns: repeat(3, 1fr);
}

.button,
button {
  background: #0f766e;
  border: 0;
  border-radius: 6px;
  color: #ffffff;
  cursor: pointer;
  display: inline-block;
  font: inherit;
  font-weight: 700;
  margin-top: 22px;
  padding: 12px 16px;
  text-align: center;
  text-decoration: none;
}

.button.secondary {
  background: #334e68;
}

.button.danger {
  background: #b42318;
}

.button:disabled,
button:disabled {
  background: #9aa6b2;
  cursor: default;
}

.actions {
  display: flex;
  flex-wrap: wrap;
  gap: 10px;
}

.actions.inline {
  align-items: center;
  justify-content: flex-end;
  margin-left: auto;
}

.button.small {
  margin: 0 6px 0 0;
  padding: 7px 10px;
}

.file-cell {
  max-width: 260px;
  overflow-wrap: anywhere;
}

td strong {
  display: block;
  margin-bottom: 3px;
}

td span {
  color: #52606d;
  display: block;
  font-size: 13px;
}

.progress {
  background: #d9e2ec;
  border-radius: 999px;
  height: 10px;
  min-width: 120px;
  overflow: hidden;
}

.progress.large {
  height: 14px;
  margin-top: 16px;
  width: 100%;
}

.progress-fill {
  background: #0f766e;
  height: 100%;
  transition: width 0.25s ease;
}

.active-job.working .progress-fill {
  background-image: linear-gradient(45deg, rgba(255, 255, 255, 0.35) 25%, transparent 25%, transparent 50%, rgba(255, 255, 255, 0.35) 50%, rgba(255, 255, 255, 0.35) 75%, transparent 75%, transparent);
  background-size: 22px 22px;
  animation: progress-stripes 0.9s linear infinite;
}

.active-job.waiting-progress .progress-fill {
  animation: progress-stripes 0.9s linear infinite, progress-waiting 1.4s ease-in-out infinite alternate;
  transition: none;
}

.progress-text {
  margin-top: 5px;
}

.active-job {
  border: 1px solid #d9e2ec;
  border-radius: 8px;
  margin: 22px 0;
  padding: 18px;
}

.active-job h2 {
  font-size: 20px;
  margin: 4px 0 6px;
  overflow-wrap: anywhere;
}

.active-job p {
  margin-bottom: 0;
}

.active-footer {
  align-items: center;
  display: flex;
  gap: 12px;
  margin-top: 12px;
}

.eyebrow {
  color: #0f766e;
  font-size: 12px;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

.queue-section {
  border-top: 1px solid #d9e2ec;
  padding: 14px 0;
}

.queue-section summary {
  cursor: pointer;
  font-weight: 700;
}

.job-list {
  list-style: none;
  margin: 12px 0 0;
  padding: 0;
}

.job-row {
  align-items: center;
  border-bottom: 1px solid #eef2f6;
  display: flex;
  gap: 14px;
  justify-content: space-between;
  padding: 10px 0;
}

.job-row span {
  overflow-wrap: anywhere;
}

.job-row small {
  color: #52606d;
  display: block;
  font-size: 12px;
  margin-top: 3px;
}

.download-cell {
  flex-shrink: 0;
}

.empty-row {
  color: #52606d;
  padding: 10px 0;
}

@keyframes progress-stripes {
  from {
    background-position: 0 0;
  }
  to {
    background-position: 22px 0;
  }
}

@keyframes progress-waiting {
  from {
    transform: translateX(0);
  }
  to {
    transform: translateX(450%);
  }
}

.back {
  color: #0f766e;
  display: inline-block;
  margin-top: 20px;
}

code {
  background: #f0f4f8;
  border-radius: 4px;
  padding: 2px 5px;
}

@media (max-width: 560px) {
  .row {
    grid-template-columns: 1fr;
  }
}
"""


HTML_PAGE = f"""
<!doctype html>
<html lang="es">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Whispermax</title>
    <style>{CSS}</style>
  </head>
  <body>
    <main class="panel">
      <h1>Whispermax</h1>
      <p>Sube uno o varios videos y Whisper generara las transcripciones en DOCX dentro de <code>salidas/transcripciones</code>.</p>
      <form method="post" action="/transcribir" enctype="multipart/form-data">
        <label for="videos">Archivos de video</label>
        <input id="videos" name="videos" type="file" accept="video/*" multiple required>

        <div class="row">
          <div>
            <label for="modelo">Modelo</label>
            <select id="modelo" name="modelo">
              <option value="tiny" selected>tiny</option>
              <option value="base">base</option>
              <option value="small">small</option>
              <option value="medium">medium</option>
              <option value="large">large</option>
            </select>
          </div>
          <div>
            <label for="idioma">Idioma</label>
            <select id="idioma" name="idioma">
              <option value="es" selected>espanol</option>
              <option value="">detectar</option>
              <option value="en">ingles</option>
              <option value="fr">frances</option>
              <option value="de">aleman</option>
              <option value="it">italiano</option>
              <option value="pt">portugues</option>
            </select>
          </div>
          <div>
            <label for="consumo">Consumo</label>
            <select id="consumo" name="consumo">
              <option value="low" selected>Bajo</option>
              <option value="balanced">Medio</option>
              <option value="fast">Rapido</option>
              <option value="ultrafast">Ultrarrapido</option>
            </select>
          </div>
        </div>

        <div class="option-row">
          <input id="marcas_tiempo" name="marcas_tiempo" type="checkbox" value="on">
          <label for="marcas_tiempo">Incluir marcas de tiempo</label>
        </div>

        <button type="submit">Anadir a cola</button>
      </form>
    </main>
  </body>
</html>
"""


def open_browser_when_ready(url: str) -> None:
    def _open() -> None:
        time.sleep(1.2)
        webbrowser.open(url)

    threading.Thread(target=_open, daemon=True).start()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Servidor local de Whispermax")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", default=8000, type=int)
    parser.add_argument("--no-browser", action="store_true")
    args = parser.parse_args()

    ensure_output_folders()
    url = f"http://{args.host}:{args.port}"
    if not args.no_browser:
        open_browser_when_ready(url)
    uvicorn.run("main:app", host=args.host, port=args.port, reload=False)
